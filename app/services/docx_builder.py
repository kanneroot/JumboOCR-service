"""
Native Markdown to DOCX Document Builder
Converts extracted Markdown text from Qwen3.7-Flash into formatted Microsoft Word (.docx) documents.
Uses pure Python standard library regex parsing and python-docx without external OS CLI dependencies.
"""

import io
import re
import logging
from typing import List, Dict, Any, Optional

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

logger = logging.getLogger(__name__)


def set_cell_background(cell, hex_color: str):
    """Sets background shading color for a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding (margins) for a table cell in dxa (1 pt = 20 dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


class DocxBuilder:
    """
    Synthesizes Markdown into a styled Microsoft Word (.docx) document.
    """

    def __init__(self):
        pass

    def build_docx(self, pages_data: List[Dict[str, Any]], doc_title: Optional[str] = None) -> io.BytesIO:
        """
        Builds a complete Word document from a list of page OCR outputs.

        :param pages_data: List of dicts containing 'page_num', 'markdown', etc.
        :param doc_title: Optional document title for metadata/header.
        :return: in-memory io.BytesIO stream of the .docx file.
        """
        doc = docx.Document()

        # Set standard 1-inch margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # Set base Normal style font
        style_normal = doc.styles['Normal']
        font = style_normal.font
        font.name = 'Calibri'
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        style_normal.paragraph_format.line_spacing = 1.15
        style_normal.paragraph_format.space_after = Pt(4)

        total_pages = len(pages_data)

        for idx, page in enumerate(pages_data):
            page_num = page.get("page_num", idx + 1)
            markdown_text = page.get("markdown", "")

            # If not the first page, add a clean page break
            if idx > 0:
                doc.add_page_break()

            # Parse and render markdown elements for this page
            self._render_page_markdown(doc, markdown_text)

        # Save document to in-memory byte buffer
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        return docx_buffer

    def _render_page_markdown(self, doc: docx.Document, markdown_text: str):
        """
        Parses Markdown text block-by-block and appends elements to doc.
        """
        if not markdown_text:
            return

        lines = markdown_text.splitlines()
        i = 0
        n = len(lines)

        while i < n:
            line = lines[i]
            stripped = line.strip()

            # 1. Blank line
            if not stripped:
                i += 1
                continue

            # 2. Horizontal Rule (---, ***, ___)
            if re.match(r'^(?:-{3,}|\*{3,}|_{3,})$', stripped):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                # Render horizontal rule using paragraph border
                pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>')
                p._p.get_or_add_pPr().append(pBdr)
                i += 1
                continue

            # 3. Headings (# Heading, ## Heading, etc.)
            heading_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                # Cap heading levels in Word (1 to 4)
                h_level = min(level, 4)
                h = doc.add_heading(level=h_level)
                self._add_formatted_runs(h, text)
                h.paragraph_format.space_before = Pt(12 if h_level == 1 else 8)
                h.paragraph_format.space_after = Pt(4)
                i += 1
                continue

            # 4. Markdown Table (| col1 | col2 |)
            if stripped.startswith('|') and stripped.endswith('|'):
                table_lines = []
                while i < n and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                self._render_markdown_table(doc, table_lines)
                continue

            # 5. Bullet List (- item, * item, + item)
            bullet_match = re.match(r'^[-*+]\s+(.*)$', stripped)
            if bullet_match:
                p = doc.add_paragraph(style='List Bullet')
                self._add_formatted_runs(p, bullet_match.group(1))
                p.paragraph_format.space_after = Pt(2)
                i += 1
                continue

            # 6. Numbered List (1. item, 2. item)
            number_match = re.match(r'^\d+\.\s+(.*)$', stripped)
            if number_match:
                p = doc.add_paragraph(style='List Number')
                self._add_formatted_runs(p, number_match.group(1))
                p.paragraph_format.space_after = Pt(2)
                i += 1
                continue

            # 7. Blockquote (> quote)
            if stripped.startswith('>'):
                quote_text = re.sub(r'^>\s*', '', stripped)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                self._add_formatted_runs(p, quote_text, default_italic=True)
                i += 1
                continue

            # 8. Standard Paragraph (may accumulate multi-line paragraph)
            para_lines = [stripped]
            i += 1
            while i < n:
                next_line = lines[i].strip()
                if not next_line:
                    break
                # Stop if next line is a special markdown element
                if (next_line.startswith('#') or
                    next_line.startswith('|') or
                    next_line.startswith('>') or
                    re.match(r'^[-*+]\s+', next_line) or
                    re.match(r'^\d+\.\s+', next_line) or
                    re.match(r'^(?:-{3,}|\*{3,}|_{3,})$', next_line)):
                    break
                para_lines.append(next_line)
                i += 1

            combined_para = " ".join(para_lines)
            p = doc.add_paragraph()
            self._add_formatted_runs(p, combined_para)

    def _add_formatted_runs(self, paragraph, text: str, default_italic: bool = False):
        """
        Parses inline markdown tokens (**bold**, *italic*, `code`) and adds styled runs to paragraph.
        """
        # Tokenizer regex: matches ***bold-italic***, **bold**, *italic*, `code`
        # Known limitation: Literal * or _ characters inside styled spans (e.g. **bold*text**)
        # and nested inline styles (e.g. **bold and *italic* inside**) are not supported.
        pattern = re.compile(
            r'(\*\*\*[^*]+\*\*\*|___[^_]+___|\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_|`[^`]+`)'
        )

        tokens = pattern.split(text)
        for token in tokens:
            if not token:
                continue

            # Bold + Italic: ***text***
            if (token.startswith('***') and token.endswith('***')) or (token.startswith('___') and token.endswith('___')):
                content = token[3:-3]
                run = paragraph.add_run(content)
                run.bold = True
                run.italic = True

            # Bold: **text** or __text__
            elif (token.startswith('**') and token.endswith('**')) or (token.startswith('__') and token.endswith('__')):
                content = token[2:-2]
                run = paragraph.add_run(content)
                run.bold = True
                if default_italic:
                    run.italic = True

            # Italic: *text* or _text_
            elif (token.startswith('*') and token.endswith('*')) or (token.startswith('_') and token.endswith('_')):
                content = token[1:-1]
                run = paragraph.add_run(content)
                run.italic = True

            # Inline code: `text`
            elif token.startswith('`') and token.endswith('`'):
                content = token[1:-1]
                run = paragraph.add_run(content)
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x99, 0x11, 0x11)

            # Plain text
            else:
                run = paragraph.add_run(token)
                if default_italic:
                    run.italic = True

    def _render_markdown_table(self, doc: docx.Document, table_lines: List[str]):
        """
        Converts a list of Markdown table lines into a styled Word table.
        """
        if not table_lines:
            return

        # Parse rows and cells
        parsed_rows = []
        for line in table_lines:
            # Strip leading and trailing pipe
            content = line.strip().strip('|')
            cells = [c.strip() for c in content.split('|')]
            parsed_rows.append(cells)

        if not parsed_rows:
            return

        # Check if row 1 (index 1) is a separator row (| --- | :---: | ---: |)
        has_header = False
        if len(parsed_rows) >= 2:
            second_row = parsed_rows[1]
            if all(re.match(r'^:?-+:?$', c) for c in second_row if c):
                has_header = True
                # Remove separator row from data
                parsed_rows.pop(1)

        # Normalize column counts
        max_cols = max(len(row) for row in parsed_rows) if parsed_rows else 1
        num_rows = len(parsed_rows)

        if num_rows == 0 or max_cols == 0:
            return

        # Create Word Table
        table = doc.add_table(rows=num_rows, cols=max_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # Set Table XML Borders for clean appearance
        tblPr = table._element.xpath('w:tblPr')
        if tblPr:
            borders = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>'
                f'<w:top w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
                f'<w:left w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
                f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
                f'<w:right w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
                f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="EAECF0"/>'
                f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="EAECF0"/>'
                f'</w:tblBorders>'
            )
            tblPr[0].append(borders)

        for row_idx, row_data in enumerate(parsed_rows):
            word_row = table.rows[row_idx]
            is_header_row = (row_idx == 0 and has_header)

            for col_idx in range(max_cols):
                cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
                cell = word_row.cells[col_idx]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell, top=120, bottom=120, left=160, right=160)

                if is_header_row:
                    set_cell_background(cell, "F2F4F7")  # Header light grey background
                elif row_idx % 2 == 1:
                    set_cell_background(cell, "F9FAFB")  # Subtle zebra row shading

                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.0

                self._add_formatted_runs(p, cell_text)

                if is_header_row:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0x10, 0x18, 0x28)

        # Add spacing after table
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(4)
        spacer.paragraph_format.space_after = Pt(4)
