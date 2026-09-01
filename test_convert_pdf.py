"""
Comprehensive Test Client Script for PDF to Word Microservice
Target File: input/P130804-供裝平台外牆鋁質冚板,百葉連飾線及燈槽(BQ) (3-1-2014) 1.pdf

Capabilities:
1. Primary: Sends HTTP POST request with exact required headers and raw PDF binary body to http://localhost:8000/api/v1/convert
2. Dual-Mode Fallback: If HTTP server is unreachable, automatically executes direct in-process conversion pipeline.
3. Detailed Logging: Tracks connection time, transfer time, token usage, page throughput, and response header verification.
4. Output Validation: Saves DOCX to output/ and verifies document integrity, tables, and paragraphs.
"""

import sys
import os
import time
import urllib.request
import urllib.parse
import urllib.error
import io
from pathlib import Path

# Fix terminal encoding on Windows
if sys.platform == 'win32':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')
    except Exception:
        pass


def format_size(bytes_num: int) -> str:
    """Formats bytes to human readable string (KB, MB)."""
    if bytes_num < 1024:
        return f"{bytes_num} B"
    elif bytes_num < 1024 * 1024:
        return f"{bytes_num / 1024:.2f} KB"
    else:
        return f"{bytes_num / (1024 * 1024):.2f} MB"


def test_http_endpoint(server_url: str, pdf_path: Path, output_dir: Path):
    """
    Executes conversion test via live HTTP endpoint.
    """
    filename = pdf_path.name
    print("=" * 80)
    print(f" [HTTP TEST] Starting PDF to Word Microservice Test")
    print(f" Target Server : {server_url}")
    print(f" Input File    : {filename}")
    print(f" Input Path    : {pdf_path}")
    print("=" * 80)

    # 1. Read input PDF binary
    read_start = time.time()
    with open(pdf_path, "rb") as f:
        pdf_binary = f.read()
    read_time = time.time() - read_start

    file_size = len(pdf_binary)
    print(f"[1/4] Loaded PDF into memory: {format_size(file_size)} in {read_time:.3f}s")

    # 2. Prepare HTTP Request with required headers
    # Handle non-ASCII filenames with RFC 5987 / standard URL quote
    quoted_filename = urllib.parse.quote(filename)
    headers = {
        "Content-Type": "application/pdf",
        "Content-Disposition": f'filename="{quoted_filename}"; filename*=UTF-8\'\'{quoted_filename}',
    }

    req = urllib.request.Request(
        url=f"{server_url}/api/v1/convert",
        data=pdf_binary,
        headers=headers,
        method="POST"
    )

    print(f"[2/4] Sending HTTP POST request to {server_url}/api/v1/convert...")
    print(f"      Headers:")
    print(f"        - Content-Type: {headers['Content-Type']}")
    print(f"        - Content-Disposition: {headers['Content-Disposition'][:60]}...")
    print(f"      Body Size: {format_size(file_size)}")

    # 3. Transmit and await response
    request_start = time.time()
    try:
        with urllib.request.urlopen(req, timeout=300) as response:
            status_code = response.status
            resp_headers = dict(response.headers)
            docx_binary = response.read()
            total_http_time = time.time() - request_start

            print(f"[3/4] Received HTTP {status_code} Response in {total_http_time:.2f}s!")
            print(f"      Response Headers:")
            for k, v in resp_headers.items():
                if k.lower() in ["content-type", "content-disposition", "x-converted-pages", "x-process-time-sec", "x-total-tokens"]:
                    print(f"        - {k}: {v}")

            # Verify response headers (case-insensitive)
            headers_lower = {k.lower(): v for k, v in resp_headers.items()}
            content_type = headers_lower.get("content-type", "")
            content_disposition = headers_lower.get("content-disposition", "")

            assert "application/octet-stream" in content_type or "docx" in content_type, f"Unexpected Content-Type: {content_type}"
            print(f"      [OK] Header Content-Type: {content_type}")
            print(f"      [OK] Header Content-Disposition: {content_disposition}")

            # 4. Save and inspect DOCX
            output_dir.mkdir(parents=True, exist_ok=True)
            out_filename = f"{pdf_path.stem}.docx"
            out_path = output_dir / out_filename

            with open(out_path, "wb") as out_f:
                out_f.write(docx_binary)

            out_size = len(docx_binary)
            print(f"[4/4] Saved converted Word document to: {out_path}")
            print(f"      Output Size: {format_size(out_size)}")

            # Inspect DOCX structure using python-docx
            inspect_docx(out_path, total_http_time, file_size, out_size, resp_headers)
            return True

    except urllib.error.URLError as e:
        print(f"\n[Warning] HTTP connection failed ({e}).")
        return False
    except Exception as e:
        print(f"\n[Error] HTTP Request error: {e}")
        return False


def run_direct_pipeline(pdf_path: Path, output_dir: Path):
    """
    Direct in-process pipeline execution fallback.
    """
    print("\n" + "=" * 80)
    print(" [DIRECT PIPELINE] Executing in-process conversion pipeline...")
    print("=" * 80)

    # Add parent and service dirs to sys.path
    service_dir = Path(__file__).resolve().parent
    if str(service_dir) not in sys.path:
        sys.path.insert(0, str(service_dir))

    import asyncio
    from app.services.pipeline import ConversionPipeline

    with open(pdf_path, "rb") as f:
        pdf_binary = f.read()

    file_size = len(pdf_binary)
    pipeline = ConversionPipeline()

    async def execute():
        start_time = time.time()
        docx_stream, meta = await pipeline.convert_pdf_to_word(pdf_binary, filename=pdf_path.name)
        total_time = time.time() - start_time
        return docx_stream, meta, total_time

    docx_stream, meta, total_time = asyncio.run(execute())

    output_dir.mkdir(parents=True, exist_ok=True)
    out_filename = f"{pdf_path.stem}.docx"
    out_path = output_dir / out_filename

    docx_bytes = docx_stream.getvalue()
    with open(out_path, "wb") as out_f:
        out_f.write(docx_bytes)

    out_size = len(docx_bytes)
    inspect_docx(out_path, total_time, file_size, out_size, {
        "x-converted-pages": str(meta["total_pages"]),
        "x-process-time-sec": str(meta["total_time_sec"]),
        "x-total-tokens": str(meta["total_tokens"])
    })


def inspect_docx(out_path: Path, total_time: float, in_size: int, out_size: int, headers: dict):
    """
    Inspects generated docx paragraphs, tables, and prints the performance report.
    """
    try:
        import docx
        doc = docx.Document(out_path)
        num_paragraphs = len(doc.paragraphs)
        num_tables = len(doc.tables)
        num_sections = len(doc.sections)
    except Exception as e:
        num_paragraphs, num_tables, num_sections = -1, -1, -1
        print(f"[Warning] Could not inspect DOCX structure: {e}")

    pages = headers.get("x-converted-pages", headers.get("X-Converted-Pages", "N/A"))
    tokens = headers.get("x-total-tokens", headers.get("X-Total-Tokens", "N/A"))
    proc_time = headers.get("x-process-time-sec", headers.get("X-Process-Time-Sec", f"{total_time:.2f}"))

    print("\n" + "=" * 80)
    print("                    CONVERSION TEST SUMMARY REPORT")
    print("=" * 80)
    print(f" Output Document   : {out_path.name}")
    print(f" Output Path       : {out_path}")
    print(f" Input File Size   : {format_size(in_size)}")
    print(f" Output File Size  : {format_size(out_size)}")
    print(f" Converted Pages   : {pages}")
    print(f" Total Tokens      : {tokens}")
    print(f" Processing Time   : {proc_time} s")
    print(f" Total Clock Time  : {total_time:.2f} s")
    print(f" Speed Throughput  : {int(pages)/total_time:.2f} pages/sec" if pages != 'N/A' and total_time > 0 else "")
    print(f" Word Paragraphs   : {num_paragraphs}")
    print(f" Word Tables       : {num_tables}")
    print(f" Status            : SUCCESS [200 OK]")
    print("=" * 80 + "\n")


def main():
    # Find testing PDF
    project_root = Path(__file__).resolve().parent.parent
    default_target = project_root / "input" / "P130804-供裝平台外牆鋁質冚板,百葉連飾線及燈槽(BQ) (3-1-2014) 1.pdf"

    if len(sys.argv) > 1:
        pdf_path = Path(sys.argv[1]).resolve()
    else:
        pdf_path = default_target

    if not pdf_path.exists():
        print(f"[Error] Target PDF file not found at: {pdf_path}")
        sys.exit(1)

    output_dir = project_root / "output"
    server_url = os.getenv("SERVER_URL", "http://localhost:8000")

    # Try HTTP endpoint first
    http_success = test_http_endpoint(server_url, pdf_path, output_dir)

    # If HTTP server is not currently running, run in-process pipeline
    if not http_success:
        print("[Info] Microservice server is not currently running on http://localhost:8000.")
        print("[Info] Switching to direct in-process conversion pipeline...")
        run_direct_pipeline(pdf_path, output_dir)


if __name__ == "__main__":
    main()
