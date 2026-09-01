"""
Comprehensive Test Suite for PDF to Word Microservice.
Tests PDF rendering, Docx synthesis, route handlers, and header extraction.
"""

import io
import pytest
import pymupdf
import docx
from unittest.mock import patch, AsyncMock
from fastapi.testclient import TestClient

from app.main import app
from app.routes import extract_filename_from_header
from app.services.pdf_renderer import PDFRenderer, PDFRenderError
from app.services.docx_builder import DocxBuilder


def create_sample_pdf_bytes() -> bytes:
    """Helper to generate a clean 2-page test PDF in memory."""
    doc = pymupdf.open()
    
    # Page 1
    page1 = doc.new_page()
    page1.insert_text((50, 72), "Document Title - Test PDF Page 1", fontsize=18)
    page1.insert_text((50, 120), "Sample paragraph content on page 1.", fontsize=12)
    
    # Page 2
    page2 = doc.new_page()
    page2.insert_text((50, 72), "Second Page Header", fontsize=16)
    page2.insert_text((50, 120), "Content on page 2.", fontsize=12)

    pdf_buffer = io.BytesIO()
    doc.save(pdf_buffer)
    doc.close()
    return pdf_buffer.getvalue()


@pytest.fixture
def client():
    with TestClient(app) as c:
        yield c


def test_health_check(client):
    """Test health check probe."""
    response = client.get("/health")
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "healthy"
    assert data["service"] == "pdf-to-word-microservice"
    assert "model" in data


def test_extract_filename_from_header():
    """Test standard and RFC 5987 filename parsing."""
    assert extract_filename_from_header('attachment; filename="annual_report.pdf"') == "annual_report.pdf"
    assert extract_filename_from_header('filename="invoice_2026.pdf"') == "invoice_2026.pdf"
    assert extract_filename_from_header('filename=statement.pdf') == "statement.pdf"
    assert extract_filename_from_header("filename*=UTF-8''quarterly%20result.pdf") == "quarterly result.pdf"
    assert extract_filename_from_header(None) == "document.pdf"


def test_pdf_renderer_valid():
    """Test in-memory PDF rendering with valid PDF bytes."""
    pdf_bytes = create_sample_pdf_bytes()
    renderer = PDFRenderer(dpi=150)
    pages = renderer.validate_and_render(pdf_bytes)

    assert len(pages) == 2
    assert pages[0]["page_num"] == 1
    assert pages[0]["total_pages"] == 2
    assert len(pages[0]["png_bytes"]) > 0
    assert pages[1]["page_num"] == 2


def test_pdf_renderer_invalid():
    """Test error handling for corrupt or invalid PDF payloads."""
    renderer = PDFRenderer(dpi=150)

    # Empty bytes
    with pytest.raises(PDFRenderError, match="empty"):
        renderer.validate_and_render(b"")

    # Non-PDF invalid bytes
    with pytest.raises(PDFRenderError, match="Invalid file format"):
        renderer.validate_and_render(b"NOT A REAL PDF FILE CONTENT")


def test_docx_builder_elements():
    """Test Markdown to DOCX conversion for headings, bold/italics, lists, and tables."""
    builder = DocxBuilder()

    sample_markdown = """# Main Document Header
This is a paragraph with **bold text**, *italic text*, and `inline code`.

## Section 2: Summary of Items
- Bullet item 1
- Bullet item 2

1. Numbered step 1
2. Numbered step 2

> This is a notable quote block.

| Item ID | Description | Unit Price | Qty | Total |
| :--- | :--- | :--- | :--- | :--- |
| ITM-001 | High Speed Scanner | $1,200.00 | 2 | $2,400.00 |
| ITM-002 | Maintenance Service | $350.00 | 1 | $350.00 |

---
"""

    pages_data = [
        {"page_num": 1, "markdown": sample_markdown},
        {"page_num": 2, "markdown": "## Page 2 Content\nAdditional details on the second page."}
    ]

    docx_buffer = builder.build_docx(pages_data, doc_title="TestDoc.pdf")
    assert docx_buffer is not None
    assert docx_buffer.getbuffer().nbytes > 0

    # Load back with docx to verify structure
    doc = docx.Document(docx_buffer)
    
    # Verify tables
    assert len(doc.tables) >= 1
    table = doc.tables[0]
    assert len(table.rows) == 3  # Header + 2 data rows
    assert len(table.columns) == 5
    assert "Item ID" in table.rows[0].cells[0].text
    assert "ITM-001" in table.rows[1].cells[0].text


def test_convert_endpoint_e2e_mock(client):
    """Test POST /api/v1/convert endpoint using mock VLM extractor."""
    pdf_bytes = create_sample_pdf_bytes()

    mock_vlm_return = [
        {
            "page_num": 1,
            "status": "success",
            "markdown": "# Financial Invoice\n| Desc | Amount |\n| --- | --- |\n| Cloud Server | $50.00 |",
            "prompt_tokens": 120,
            "completion_tokens": 45,
            "total_tokens": 165
        },
        {
            "page_num": 2,
            "status": "success",
            "markdown": "## Terms & Conditions\nThank you for your business.",
            "prompt_tokens": 80,
            "completion_tokens": 20,
            "total_tokens": 100
        }
    ]

    with patch("app.services.vlm_ocr.VLMExtractor.extract_document", new_callable=AsyncMock) as mock_extract:
        mock_extract.return_value = mock_vlm_return

        response = client.post(
            "/api/v1/convert",
            headers={
                "Content-Type": "application/pdf",
                "Content-Disposition": 'filename="sample_invoice.pdf"'
            },
            content=pdf_bytes
        )

        assert response.status_code == 200
        assert response.headers["content-type"] == "application/octet-stream"
        assert 'filename="sample_invoice.docx"' in response.headers["content-disposition"]
        assert response.headers["x-converted-pages"] == "2"

        # Verify returned content is a valid Word docx
        docx_stream = io.BytesIO(response.content)
        doc = docx.Document(docx_stream)
        assert len(doc.tables) >= 1
        assert "Financial Invoice" in [p.text for p in doc.paragraphs if p.text]
